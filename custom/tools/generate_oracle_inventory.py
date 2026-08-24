from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt


CUSTOM_ROOT = Path(__file__).resolve().parents[1]
FORGE_ROOT = CUSTOM_ROOT.parent
CARDS_ROOT = CUSTOM_ROOT / "cards"
EDITIONS_ROOT = CUSTOM_ROOT / "editions"
LOCALIZATION = FORGE_ROOT / "forge-gui" / "res" / "languages" / "cardnames-zh-CN.txt"
TEXT_OUTPUT = CUSTOM_ROOT / "DIY卡牌_游戏内Oracle_非测试卡.txt"
DOCX_OUTPUT = CUSTOM_ROOT / "DIY卡牌_游戏内Oracle_非测试卡.docx"

TITAN_NAMES = {
    "阿曼苏尔",
    "灭世者萨格拉斯",
    "兵主",
    "生命的缚誓者艾欧娜尔",
    "维和者阿米图斯",
    "诺甘农",
    "雷霆之神高戈奈斯",
    "卡兹格罗斯",
    "翠绿之星阿古斯",
    "复仇者阿格拉玛",
}

EDITION_ORDER = (
    "Placeholder_Set.txt",
    "JiFei99_Set.txt",
    "BoTu_Three_Kingdoms_New_Chapter.txt",
)


def active_card_names() -> set[str]:
    names = set()
    tracked = subprocess.run(
        ["git", "-c", "core.quotepath=false", "ls-files", "-z", "--", "custom/cards"],
        cwd=FORGE_ROOT,
        capture_output=True,
        text=True,
        encoding="utf-8",
        check=False,
    )
    if tracked.returncode == 0:
        paths = [
            FORGE_ROOT / relative
            for relative in tracked.stdout.split("\0")
            if relative.endswith(".txt") and (FORGE_ROOT / relative).is_file()
        ]
    else:
        paths = list(CARDS_ROOT.rglob("*.txt"))

    for path in paths:
        if path.name.lower().startswith("test") or path.name == "forge_test_goblin.txt":
            continue
        for line in path.read_text(encoding="utf-8").splitlines():
            if line.startswith("Name:"):
                names.add(line.removeprefix("Name:"))
                break
    return names


def localized_cards() -> dict[str, tuple[str, str]]:
    entries = {}
    for line in LOCALIZATION.read_text(encoding="utf-8").splitlines():
        parts = line.split("|", 3)
        if len(parts) == 4:
            entries[parts[0]] = (parts[2], parts[3].replace(r"\n", "\n"))
    return entries


def ordered_cards(names: set[str]) -> list[tuple[str, str, str]]:
    ordered = []
    seen = set()
    candidates = sorted(names, key=len, reverse=True)

    for filename in EDITION_ORDER:
        path = EDITIONS_ROOT / filename
        code = next(
            line.removeprefix("Code=")
            for line in path.read_text(encoding="utf-8").splitlines()
            if line.startswith("Code=")
        )
        for line in path.read_text(encoding="utf-8").splitlines():
            match = re.match(r"^(\d+)\s+[A-Z]+\s+(.+)$", line)
            if not match:
                continue
            collector, remainder = match.groups()
            name = next(
                (
                    candidate
                    for candidate in candidates
                    if remainder == candidate or remainder.startswith(candidate + " @")
                ),
                None,
            )
            if name is not None and name not in seen:
                ordered.append((code, collector, name))
                seen.add(name)

    for name in sorted(names - seen):
        ordered.append(("未登记", "—", name))
    return ordered


def inventory_rows() -> list[tuple[str, str, str, str, str]]:
    names = active_card_names()
    localized = localized_cards()
    missing = sorted(names - localized.keys())
    if missing:
        raise RuntimeError(f"Missing zh-CN localization: {missing}")

    rows = []
    for code, collector, name in ordered_cards(names):
        card_type, oracle = localized[name]
        rows.append((code, collector, name, card_type, oracle))
    return rows


def write_text(rows: list[tuple[str, str, str, str, str]]) -> None:
    audited = len(rows) - len(TITAN_NAMES)
    lines = [
        "Forge DIY 卡牌：游戏内真实 Oracle（非测试卡）",
        "================================================",
        "",
        "来源：custom/cards、非测试版本登记与 forge-gui/res/languages/cardnames-zh-CN.txt",
        f"口径：已隐藏 Gigantic Spright；当前共 {len(rows)} 张非测试 DIY 卡。",
        f"本次 wording 审校范围：{audited} 张；未改动 {len(TITAN_NAMES)} 张炉石泰坦卡牌。",
        "说明：以下名称、类别和 Oracle 均按当前源码生成；每次文字变更后应重新运行本工具。",
        "",
    ]

    for index, (code, collector, name, card_type, oracle) in enumerate(rows, 1):
        lines.extend(
            [
                f"{index:03d}. {name}",
                f"版本：{code} / {collector}",
                f"类别：{card_type}",
                "Oracle：" + oracle,
                "",
            ]
        )

    TEXT_OUTPUT.write_text("\r\n".join(lines), encoding="utf-8", newline="")


def set_run_font(run, name: str = "Microsoft YaHei", size: int = 10) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_multiline_paragraph(document: Document, prefix: str, value: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(prefix)
    run.bold = True
    set_run_font(run)
    for index, part in enumerate(value.splitlines()):
        if index:
            run.add_break(WD_BREAK.LINE)
        run = paragraph.add_run(part)
        set_run_font(run)


def write_docx(rows: list[tuple[str, str, str, str, str]]) -> None:
    audited = len(rows) - len(TITAN_NAMES)
    document = Document()
    styles = document.styles
    for style_name, size in (("Normal", 10), ("Title", 20), ("Heading 1", 13)):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    document.add_heading("Forge DIY 卡牌：游戏内真实 Oracle（非测试卡）", 0)
    document.add_paragraph(
        "来源：custom/cards、非测试版本登记与 cardnames-zh-CN.txt"
    )
    document.add_paragraph(
        f"已隐藏 Gigantic Spright；当前共 {len(rows)} 张非测试 DIY 卡。"
        f"本次 wording 审校 {audited} 张，未改动 {len(TITAN_NAMES)} 张炉石泰坦卡牌。"
    )

    for index, (code, collector, name, card_type, oracle) in enumerate(rows, 1):
        document.add_heading(f"{index:03d}. {name}", level=1)
        add_multiline_paragraph(document, "版本：", f"{code} / {collector}")
        add_multiline_paragraph(document, "类别：", card_type)
        add_multiline_paragraph(document, "Oracle：", oracle)

    document.save(DOCX_OUTPUT)


def main() -> None:
    rows = inventory_rows()
    write_text(rows)
    write_docx(rows)
    print(f"Wrote {len(rows)} cards to the text and DOCX inventories")


if __name__ == "__main__":
    main()
